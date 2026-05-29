from __future__ import annotations

import sqlite3
from pathlib import Path

from docx import Document
from docx.shared import Pt

from app.analytics import strategy_report
from app.forecasting import evaluate_day_ahead_model


EXPORT_DIR = Path(__file__).resolve().parents[2] / "data" / "exports"


def _fmt(value: object, digits: int = 1) -> str:
    if value is None:
        return "-"
    if isinstance(value, float):
        return f"{value:.{digits}f}"
    return str(value)


def _set_font(document: Document) -> None:
    styles = document.styles
    for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style.font.size = Pt(10.5 if style_name == "Normal" else 14)


def _add_key_value_table(document: Document, rows: list[tuple[str, str]]) -> None:
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "项目"
    table.rows[0].cells[1].text = "内容"
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value


def _add_hourly_table(document: Document, hourly_advice: list[dict]) -> None:
    table = document.add_table(rows=1, cols=8)
    table.style = "Table Grid"
    headers = ["小时", "风险", "倾向", "日前价", "实时价", "价差", "报价区间", "主要原因"]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for item in hourly_advice:
        cells = table.add_row().cells
        cells[0].text = f"{item['hour']:02d}:00"
        cells[1].text = item.get("risk_level") or "-"
        cells[2].text = item.get("stance") or "-"
        cells[3].text = _fmt(item.get("day_ahead_price"))
        cells[4].text = _fmt(item.get("real_time_price"))
        cells[5].text = _fmt(item.get("spread_real_minus_day_ahead"))
        cells[6].text = f"{_fmt(item.get('quote_lower'), 0)}-{_fmt(item.get('quote_upper'), 0)}"
        cells[7].text = "；".join((item.get("reasons") or [])[:3])


def export_strategy_docx(
    conn: sqlite3.Connection,
    market_date: str,
    real_time_method: str = "spread_follow",
) -> Path:
    report = strategy_report(conn, market_date, real_time_method=real_time_method)
    evaluation = evaluate_day_ahead_model(conn, end_date=market_date, days=3)
    summary = report["summary"]
    hourly_advice = report["hourly_advice"]

    EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    output_path = EXPORT_DIR / f"广西现货报价辅助报告_{market_date}.docx"

    document = Document()
    _set_font(document)
    document.add_heading(f"广西现货报价辅助报告（{market_date}）", level=1)
    document.add_paragraph(f"模式：{'目标日预测' if report['mode'] == 'target_day' else '复盘分析'}")
    document.add_paragraph(f"实时价格方案：{real_time_method}")

    document.add_heading("一、核心结论", level=2)
    document.add_paragraph(report["headline"])
    for item in report.get("narrative", []):
        document.add_paragraph(item, style=None)

    document.add_heading("二、价格概览", level=2)
    _add_key_value_table(
        document,
        [
            ("日前均价", f"{_fmt(summary['day_ahead']['avg'])} 元/MWh"),
            ("日前峰谷差", f"{_fmt(summary['day_ahead']['range'])} 元/MWh"),
            ("实时均价", f"{_fmt(summary['real_time']['avg'])} 元/MWh"),
            ("实时-日前平均价差", f"{_fmt(summary['spread_real_minus_day_ahead']['avg'])} 元/MWh"),
            ("高风险时段", "、".join(f"{hour}:00" for hour in report.get("high_risk_hours", [])[:10]) or "-"),
        ],
    )

    document.add_heading("三、供需边际口径", level=2)
    document.add_paragraph("供给：水电 + 新能源出力。")
    document.add_paragraph("需求：统调负荷 + 省间联络线。")
    supply = report.get("supply_series", [])
    if supply:
        peak_demand = max(supply, key=lambda row: row.get("demand_total") or 0)
        min_supply_ratio = min(
            [row for row in supply if row.get("demand_total")],
            key=lambda row: row.get("supply_demand_ratio") or 0,
        )
        _add_key_value_table(
            document,
            [
                (
                    "最大需求小时",
                    f"{peak_demand['hour']:02d}:00，需求 {_fmt(peak_demand.get('demand_total'), 0)} MW",
                ),
                (
                    "供给覆盖率最低小时",
                    f"{min_supply_ratio['hour']:02d}:00，覆盖率 {_fmt((min_supply_ratio.get('supply_demand_ratio') or 0) * 100)}%",
                ),
            ],
        )

    document.add_heading("四、模型回测摘要", level=2)
    eval_summary = evaluation.get("summary", {})
    _add_key_value_table(
        document,
        [
            ("回测日期", "、".join(evaluation.get("dates", [])) or "-"),
            ("集成模型 MAE", f"{_fmt(eval_summary.get('ensemble', {}).get('mae'))} 元/MWh"),
            ("集成模型 RMSE", f"{_fmt(eval_summary.get('ensemble', {}).get('rmse'))} 元/MWh"),
            ("XGBoost MAE", f"{_fmt(eval_summary.get('xgboost', {}).get('mae'))} 元/MWh"),
            ("LSTM MAE", f"{_fmt(eval_summary.get('lstm', {}).get('mae'))} 元/MWh"),
        ],
    )

    document.add_heading("五、小时报价建议", level=2)
    _add_hourly_table(document, hourly_advice)

    document.add_paragraph("说明：本报告基于已导入的广西原始现货数据自动生成，预测结果仅用于报价辅助和风险提示。")
    document.save(output_path)
    return output_path
